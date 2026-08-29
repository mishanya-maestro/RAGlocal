"""Извлечение текста, классификация и структурный парсинг документов."""

from __future__ import annotations

import io
import os
import re
import shutil
import tempfile
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import config
from generator import _call_llm, _clean_llm_output


@dataclass
class DocumentSegment:
    """Смысловой фрагмент документа с метаданными."""

    text: str
    page: int | None = None
    section: str | None = None
    segment_type: str = "paragraph"  # paragraph, heading, table
    index: int = 0


@dataclass
class ParsedDocument:
    """Результат парсинга загруженного документа."""

    filename: str
    doc_type: str = "unknown"
    doc_type_confidence: float = 0.0
    segments: list[DocumentSegment] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    full_text: str = ""

    @property
    def doc_type_label(self) -> str:
        return _TYPE_LABELS.get(self.doc_type, "Прочий документ")


# Ключевые слова для rule-based классификации.
_TYPE_KEYWORDS: dict[str, list[str]] = {
    "contract": [
        "договор",
        "стороны",
        "исполнитель",
        "заказчик",
        "арендодатель",
        "арендатор",
        "продавец",
        "покупатель",
        "предмет договора",
        "цена договора",
        "срок действия договора",
    ],
    "claim": [
        "исковое заявление",
        "истец",
        "ответчик",
        "предмет иска",
        "основание иска",
        "государственная пошлина",
        "в суд",
        "требования",
        "просил",
    ],
    "order": [
        "приказ",
        "распоряжение",
        "приказываю",
        "работник",
        "должностная инструкция",
        "уволить",
        "принять на работу",
        "основании",
    ],
    "labor_contract": [
        "трудовой договор",
        "работодатель",
        "работник",
        "должность",
        "заработная плата",
        "испытательный срок",
        "режим рабочего времени",
    ],
}

_TYPE_LABELS: dict[str, str] = {
    "contract": "Договор",
    "claim": "Исковое заявление",
    "order": "Приказ/распоряжение",
    "labor_contract": "Трудовой договор",
    "unknown": "Прочий документ",
}


def _extract_text_from_pdf_bytes(file_bytes: bytes) -> list[DocumentSegment]:
    """Извлекает текст из PDF с сохранением структуры страниц."""
    try:
        import fitz
    except ImportError as e:
        raise RuntimeError("Не установлен pymupdf. Выполните: pip install pymupdf") from e

    segments: list[DocumentSegment] = []
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        pdf_path = tmp.name

    try:
        doc = fitz.open(pdf_path)
        for page_idx, page in enumerate(doc, start=1):
            # Пытаемся сохранить блочную структуру (заголовки/абзацы).
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))  # top-to-bottom, left-to-right
            for idx, block in enumerate(blocks):
                text = block[4].strip()
                if not text:
                    continue
                is_heading = bool(
                    block[3] - block[1] > 14  # высокий блок
                    or re.match(r"^(?:Статья\s+\d+|Раздел|Глава|\d+\.\s+)", text, re.IGNORECASE)
                )
                segments.append(
                    DocumentSegment(
                        text=text,
                        page=page_idx,
                        segment_type="heading" if is_heading else "paragraph",
                        index=len(segments),
                    )
                )
        doc.close()
    finally:
        try:
            os.unlink(pdf_path)
        except OSError:
            pass

    return segments


def _extract_text_from_docx_bytes(file_bytes: bytes) -> list[DocumentSegment]:
    """Извлекает текст из DOCX с сохранением параграфов и таблиц."""
    try:
        import docx
    except ImportError as e:
        raise RuntimeError(
            "Не установлен python-docx. Выполните: pip install python-docx"
        ) from e

    segments: list[DocumentSegment] = []
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        docx_path = tmp.name

    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_heading = bool(
                para.style.name.startswith("Heading")
                or re.match(r"^(?:Статья\s+\d+|Раздел|Глава|\d+\.\s+)", text, re.IGNORECASE)
            )
            segments.append(
                DocumentSegment(
                    text=text,
                    segment_type="heading" if is_heading else "paragraph",
                    index=len(segments),
                )
            )
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    segments.append(
                        DocumentSegment(
                            text=" | ".join(cells),
                            segment_type="table",
                            index=len(segments),
                        )
                    )
    finally:
        try:
            os.unlink(docx_path)
        except OSError:
            pass

    return segments


def _extract_text_from_image_bytes(file_bytes: bytes) -> list[DocumentSegment]:
    """Распознаёт текст на фото/скане через Tesseract OCR."""
    from PIL import Image

    try:
        image = Image.open(io.BytesIO(file_bytes))
    except Exception as e:
        raise RuntimeError("Не удалось открыть изображение") from e

    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")

    text = _ocr_with_pytesseract(image)

    if not text:
        return []

    segments: list[DocumentSegment] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        is_heading = bool(
            re.match(r"^(?:Статья\s+\d+|Раздел|Глава|\d+\.\s+)", line, re.IGNORECASE)
            or line.isupper()
        )
        segments.append(
            DocumentSegment(
                text=line,
                segment_type="heading" if is_heading else "paragraph",
                index=len(segments),
            )
        )
    return segments


def _ocr_with_pytesseract(image) -> str:
    """Fallback на Tesseract OCR."""
    try:
        import pytesseract
    except ImportError as e:
        raise RuntimeError(
            "Не установлен OCR. Выполните:\n"
            "  pip install pytesseract\n"
            "  + установите Tesseract: winget install --id UB-Mannheim.TesseractOCR\n"
            "  + скачайте rus.traineddata в C:\\Program Files\\Tesseract-OCR\\tessdata"
        ) from e

    # Автоопределение пути к tesseract.exe.
    tesseract_path = shutil.which("tesseract")
    if not tesseract_path:
        for candidate in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.exists(candidate):
                tesseract_path = candidate
                break
    if not tesseract_path or not os.path.exists(tesseract_path):
        raise RuntimeError(
            "Tesseract OCR не найден. Установите Tesseract: "
            "winget install --id UB-Mannheim.TesseractOCR, "
            "или скачайте установщик с https://github.com/UB-Mannheim/tesseract/wiki"
        )
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    # Определяем, доступен ли русский язык.
    tessdata_dir = os.path.join(os.path.dirname(tesseract_path), "tessdata")
    lang = "rus+eng" if os.path.exists(os.path.join(tessdata_dir, "rus.traineddata")) else "eng"

    try:
        return pytesseract.image_to_string(image, lang=lang)
    except Exception as e:
        raise RuntimeError(f"Ошибка OCR: {e}") from e


def _extract_text_from_txt_bytes(file_bytes: bytes) -> list[DocumentSegment]:
    """Декодирует TXT и разбивает на параграфы/заголовки."""
    text = ""
    for enc in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue

    if not text:
        raise RuntimeError("Не удалось декодировать текстовый файл")

    segments: list[DocumentSegment] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        is_heading = bool(
            re.match(r"^(?:Статья\s+\d+|Раздел|Глава|\d+\.\s+)", line, re.IGNORECASE)
            or line.isupper()
        )
        segments.append(
            DocumentSegment(
                text=line,
                segment_type="heading" if is_heading else "paragraph",
                index=len(segments),
            )
        )
    return segments


IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif")


def extract_text(filename: str, file_bytes: bytes) -> list[DocumentSegment]:
    """Универсальное извлечение текста по расширению файла."""
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return _extract_text_from_pdf_bytes(file_bytes)
    if lower.endswith(".docx"):
        return _extract_text_from_docx_bytes(file_bytes)
    if lower.endswith(".txt"):
        return _extract_text_from_txt_bytes(file_bytes)
    if lower.endswith(IMAGE_EXTENSIONS):
        return _extract_text_from_image_bytes(file_bytes)
    raise RuntimeError(f"Неподдерживаемый формат файла: {filename}")


def _rule_based_classify(text: str) -> tuple[str, float]:
    """Классификация по ключевым словам. Возвращает тип и уверенность."""
    text_lower = text.lower()
    scores: dict[str, int] = {}
    for doc_type, keywords in _TYPE_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score:
            scores[doc_type] = score

    if not scores:
        return "unknown", 0.0

    best = max(scores, key=scores.get)
    total = sum(scores.values())
    confidence = scores[best] / total if total else 0.0
    return best, confidence


def _llm_classify(text: str) -> tuple[str, float]:
    """LLM fallback для классификации, если rule-based не уверен."""
    sample = text[:3000]
    prompt = f"""Классифицируй документ по одному из типов: Договор, Исковое заявление, Приказ/распоряжение, Трудовой договор, Прочий документ.

Текст документа (начало):
{sample}

Ответь строго в формате JSON:
{{"type": "contract|claim|order|labor_contract|unknown", "confidence": 0.0-1.0}}

JSON:"""
    try:
        raw = _call_llm(
            system_prompt="Ты юридический классификатор документов. Ответ только JSON.",
            user_prompt=prompt,
            mode_override=None,
            model_override="qwen3.5:4b",
        )
        raw = _clean_llm_output(raw)
        if "```" in raw:
            raw = raw.split("```")[-2] if raw.count("```") >= 2 else raw.split("```")[-1]
        raw = raw.strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
        import json

        data = json.loads(raw)
        doc_type = data.get("type", "unknown")
        confidence = float(data.get("confidence", 0.0))
        return doc_type, confidence
    except Exception as e:
        traceback.print_exc()
        return "unknown", 0.0


def classify_document(text: str) -> tuple[str, float]:
    """Определяет тип документа."""
    doc_type, confidence = _rule_based_classify(text)
    if doc_type == "unknown" or confidence < 0.4:
        llm_type, llm_confidence = _llm_classify(text)
        if llm_confidence > confidence:
            return llm_type, llm_confidence
    return doc_type, confidence


def _extract_metadata(text: str) -> dict[str, Any]:
    """Извлекает базовые реквизиты документа."""
    metadata: dict[str, Any] = {}

    # Даты вида 12.05.2024, 12.05.2024 г.
    dates = re.findall(r"\b\d{2}\.\d{2}\.\d{4}\s*(?:г\.?)?\b", text)
    if dates:
        metadata["dates"] = list(dict.fromkeys(dates[:5]))

    # Суммы с валютой.
    sums = re.findall(r"\b(\d{1,3}(?:\s?\d{3})*(?:,\d{2})?\s*(?:руб|USD|EUR|бел\.\s*руб))\b", text, re.IGNORECASE)
    if sums:
        metadata["amounts"] = list(dict.fromkeys(sums[:5]))

    # Стороны (эвристика).
    parties: list[str] = []
    for pattern in [
        r"(?:Исполнитель|Продавец|Арендодатель|Работодатель):\s*([^\n,]+)",
        r"(?:Заказчик|Покупатель|Арендатор|Работник):\s*([^\n,]+)",
    ]:
        for m in re.finditer(pattern, text, re.IGNORECASE):
            party = m.group(1).strip('"«» .').strip()
            if party and len(party) > 2 and party not in parties:
                parties.append(party)
    if parties:
        metadata["parties"] = parties[:6]

    return metadata


def parse_uploaded_document(filename: str, file_bytes: bytes) -> ParsedDocument:
    """Полный pipeline: извлечение, классификация, метаданные."""
    segments = extract_text(filename, file_bytes)
    full_text = "\n\n".join(s.text for s in segments)

    doc_type, confidence = classify_document(full_text)
    metadata = _extract_metadata(full_text)
    metadata["segment_count"] = len(segments)
    metadata["char_count"] = len(full_text)

    return ParsedDocument(
        filename=filename,
        doc_type=doc_type,
        doc_type_confidence=confidence,
        segments=segments,
        metadata=metadata,
        full_text=full_text,
    )


def chunk_segments(
    segments: list[DocumentSegment], max_chunk_size: int = 1500, overlap: int = 200
) -> list[DocumentSegment]:
    """Склеивает сегменты в чанки с overlap, сохраняя метаданные первого сегмента."""
    chunks: list[DocumentSegment] = []
    current_text = ""
    current_meta: dict[str, Any] = {}

    for seg in segments:
        candidate = (current_text + "\n\n" + seg.text).strip() if current_text else seg.text
        if len(candidate) <= max_chunk_size:
            current_text = candidate
            if not current_meta:
                current_meta = {
                    "page": seg.page,
                    "section": seg.section,
                    "segment_type": seg.segment_type,
                    "start_index": seg.index,
                }
        else:
            if current_text:
                chunks.append(
                    DocumentSegment(
                        text=current_text,
                        page=current_meta.get("page"),
                        section=current_meta.get("section"),
                        segment_type=current_meta.get("segment_type", "paragraph"),
                        index=len(chunks),
                    )
                )
            current_text = seg.text
            current_meta = {
                "page": seg.page,
                "section": seg.section,
                "segment_type": seg.segment_type,
                "start_index": seg.index,
            }

    if current_text:
        chunks.append(
            DocumentSegment(
                text=current_text,
                page=current_meta.get("page"),
                section=current_meta.get("section"),
                segment_type=current_meta.get("segment_type", "paragraph"),
                index=len(chunks),
            )
        )

    # overlap: добавить хвост предыдущего чанка в начало следующего.
    if overlap > 0 and len(chunks) > 1:
        overlapped: list[DocumentSegment] = []
        for i, ch in enumerate(chunks):
            prefix = ""
            if i > 0:
                prev = chunks[i - 1].text
                if len(prev) > overlap:
                    prefix = prev[-overlap:].lstrip()
                else:
                    prefix = prev
            text = (prefix + "\n\n" + ch.text).strip() if prefix else ch.text
            overlapped.append(
                DocumentSegment(
                    text=text,
                    page=ch.page,
                    section=ch.section,
                    segment_type=ch.segment_type,
                    index=i,
                )
            )
        chunks = overlapped

    return chunks
